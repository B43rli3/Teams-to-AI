"""Verarbeitung von Teams-Anhängen (Bilder und Dokumente)."""

from __future__ import annotations

import base64
import io
import re
from dataclasses import dataclass, field
from typing import Any
from urllib.parse import unquote

from app.config import Settings
from app.graph_client import GraphClient
from app.logging_config import get_logger, truncate_id, truncate_text

logger = get_logger(__name__)

IMAGE_EXTENSIONS = frozenset({".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"})
IMAGE_CONTENT_TYPES = frozenset(
    {
        "image/png",
        "image/jpeg",
        "image/jpg",
        "image/gif",
        "image/webp",
        "image/bmp",
    }
)
DOCUMENT_EXTENSIONS = frozenset(
    {".pdf", ".docx", ".txt", ".md", ".csv", ".json", ".xml", ".log", ".html", ".htm"}
)

HOSTED_CONTENT_RE = re.compile(
    r"hostedContents/([^/\"'\s]+)/\$value",
    re.IGNORECASE,
)


@dataclass
class ProcessedAttachment:
    """Ergebnis der Verarbeitung eines einzelnen Anhangs."""

    name: str
    kind: str  # image | document | unsupported
    text_excerpt: str | None = None
    image_base64: str | None = None
    notes: str | None = None


@dataclass
class AttachmentBundle:
    """Zusammengefasste Anhangsdaten für den LLM-Aufruf."""

    images_base64: list[str] = field(default_factory=list)
    document_texts: list[str] = field(default_factory=list)
    notes: list[str] = field(default_factory=list)
    processed: list[ProcessedAttachment] = field(default_factory=list)

    @property
    def has_content(self) -> bool:
        return bool(self.images_base64 or self.document_texts)

    def build_context_block(self) -> str:
        """Baut einen Textblock mit Dokumentinhalten und Hinweisen."""
        parts: list[str] = []
        for text in self.document_texts:
            parts.append(text)
        for note in self.notes:
            parts.append(note)
        if self.images_base64:
            parts.append(
                f"[Es wurden {len(self.images_base64)} Bild(er) als visuelle Eingabe "
                "an das Modell übergeben.]"
            )
        return "\n\n".join(parts).strip()


class AttachmentProcessor:
    """Lädt Teams-Anhänge herunter und bereitet sie für das LLM vor."""

    def __init__(self, graph_client: GraphClient, settings: Settings) -> None:
        self._graph = graph_client
        self._settings = settings

    async def process_message_attachments(
        self,
        message_id: str,
        body_html: str,
        attachments: list[dict[str, Any]],
    ) -> AttachmentBundle:
        """Verarbeitet Anhänge und Inline-Bilder einer Nachricht."""
        if not self._settings.process_attachments:
            return AttachmentBundle()

        bundle = AttachmentBundle()
        seen_ids: set[str] = set()
        max_files = self._settings.attachment_max_files
        processed_count = 0

        hosted_ids = self._extract_hosted_content_ids(body_html)
        for hosted_id in hosted_ids:
            if processed_count >= max_files:
                bundle.notes.append(
                    f"[Weitere Anhänge übersprungen (Limit {max_files}).]"
                )
                break
            if hosted_id in seen_ids:
                continue
            seen_ids.add(hosted_id)
            item = await self._process_hosted_content(message_id, hosted_id)
            self._merge_item(bundle, item)
            processed_count += 1

        for raw in attachments:
            if processed_count >= max_files:
                bundle.notes.append(
                    f"[Weitere Anhänge übersprungen (Limit {max_files}).]"
                )
                break
            att_id = str(raw.get("id") or raw.get("contentUrl") or raw.get("name") or "")
            if att_id and att_id in seen_ids:
                continue
            if att_id:
                seen_ids.add(att_id)
            item = await self._process_attachment_meta(message_id, raw)
            self._merge_item(bundle, item)
            processed_count += 1

        logger.info(
            "attachments_processed",
            message_id=truncate_id(message_id),
            images=len(bundle.images_base64),
            documents=len(bundle.document_texts),
            notes=len(bundle.notes),
        )
        return bundle

    def _merge_item(self, bundle: AttachmentBundle, item: ProcessedAttachment) -> None:
        bundle.processed.append(item)
        if item.image_base64:
            bundle.images_base64.append(item.image_base64)
        if item.text_excerpt:
            bundle.document_texts.append(
                f"--- Dokument: {item.name} ---\n{item.text_excerpt}"
            )
        if item.notes:
            bundle.notes.append(item.notes)

    async def _process_hosted_content(
        self, message_id: str, hosted_id: str
    ) -> ProcessedAttachment:
        name = f"hosted-{truncate_id(hosted_id, 12)}"
        try:
            data, content_type = await self._graph.download_hosted_content(
                message_id=message_id,
                hosted_content_id=hosted_id,
                team_id=self._settings.teams_team_id or None,
                channel_id=self._settings.teams_channel_id or None,
                chat_id=self._settings.teams_chat_id or None,
                target_mode=self._settings.teams_target_mode,
            )
        except Exception as exc:
            logger.warning(
                "hosted_content_download_failed",
                hosted_id=truncate_id(hosted_id),
                error=str(exc)[:200],
            )
            return ProcessedAttachment(
                name=name,
                kind="unsupported",
                notes=f"[Inline-Inhalt konnte nicht geladen werden: {name}]",
            )

        return self._classify_and_extract(name, data, content_type)

    async def _process_attachment_meta(
        self, message_id: str, raw: dict[str, Any]
    ) -> ProcessedAttachment:
        name = str(raw.get("name") or "anhang")
        content_type = str(raw.get("contentType") or "").lower()
        content_url = raw.get("contentUrl")
        content_bytes_b64 = raw.get("contentBytes")

        data: bytes | None = None
        resolved_type = content_type

        if isinstance(content_bytes_b64, str) and content_bytes_b64:
            try:
                data = base64.b64decode(content_bytes_b64)
            except Exception:
                data = None

        if data is None and content_url:
            try:
                data, resolved_type = await self._graph.download_binary_url(
                    str(content_url)
                )
            except Exception as exc:
                logger.warning(
                    "attachment_url_download_failed",
                    name=truncate_text(name, 40),
                    error=str(exc)[:200],
                )
                return ProcessedAttachment(
                    name=name,
                    kind="unsupported",
                    notes=(
                        f"[Anhang '{name}' konnte nicht heruntergeladen werden. "
                        "Datei liegt ggf. in SharePoint und benötigt zusätzliche Graph-Rechte.]"
                    ),
                )

        # Teams speichert Bild-Upload oft nur als Reference ohne contentUrl.
        # Dann Hosted Contents nachladen.
        if data is None and "image" in content_type:
            hosted_ids = await self._graph.list_hosted_content_ids(
                message_id=message_id,
                team_id=self._settings.teams_team_id or None,
                channel_id=self._settings.teams_channel_id or None,
                chat_id=self._settings.teams_chat_id or None,
                target_mode=self._settings.teams_target_mode,
            )
            for hosted_id in hosted_ids:
                try:
                    data, resolved_type = await self._graph.download_hosted_content(
                        message_id=message_id,
                        hosted_content_id=hosted_id,
                        team_id=self._settings.teams_team_id or None,
                        channel_id=self._settings.teams_channel_id or None,
                        chat_id=self._settings.teams_chat_id or None,
                        target_mode=self._settings.teams_target_mode,
                    )
                    break
                except Exception:
                    continue

        if data is None:
            return ProcessedAttachment(
                name=name,
                kind="unsupported",
                notes=f"[Anhang '{name}' konnte nicht gelesen werden.]",
            )

        return self._classify_and_extract(name, data, resolved_type)

    def _classify_and_extract(
        self, name: str, data: bytes, content_type: str
    ) -> ProcessedAttachment:
        max_bytes = self._settings.attachment_max_bytes
        if len(data) > max_bytes:
            return ProcessedAttachment(
                name=name,
                kind="unsupported",
                notes=(
                    f"[Anhang '{name}' zu groß ({len(data)} Bytes, "
                    f"Limit {max_bytes}).]"
                ),
            )

        lower_name = name.lower()
        ext = ""
        if "." in lower_name:
            ext = "." + lower_name.rsplit(".", 1)[-1]

        ctype = (content_type or "").split(";")[0].strip().lower()

        is_image = (
            ext in IMAGE_EXTENSIONS
            or ctype in IMAGE_CONTENT_TYPES
            or ctype.startswith("image/")
        )
        is_document = ext in DOCUMENT_EXTENSIONS or ctype in {
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "text/plain",
            "text/markdown",
            "text/csv",
            "application/json",
            "text/html",
            "application/xml",
            "text/xml",
        }

        if is_image:
            if not self._settings.process_images:
                return ProcessedAttachment(
                    name=name,
                    kind="unsupported",
                    notes=f"[Bild '{name}' ignoriert (PROCESS_IMAGES=false).]",
                )
            encoded = base64.b64encode(data).decode("ascii")
            return ProcessedAttachment(
                name=name,
                kind="image",
                image_base64=encoded,
            )

        if is_document:
            if not self._settings.process_documents:
                return ProcessedAttachment(
                    name=name,
                    kind="unsupported",
                    notes=f"[Dokument '{name}' ignoriert (PROCESS_DOCUMENTS=false).]",
                )
            text = extract_document_text(name, data, ctype)
            if not text:
                return ProcessedAttachment(
                    name=name,
                    kind="unsupported",
                    notes=f"[Dokument '{name}' enthielt keinen extrahierbaren Text.]",
                )
            max_chars = self._settings.attachment_max_document_chars
            truncated = False
            if len(text) > max_chars:
                text = text[:max_chars]
                truncated = True
            if truncated:
                text += "\n\n[Dokumenttext gekürzt.]"
            return ProcessedAttachment(
                name=name,
                kind="document",
                text_excerpt=text,
            )

        return ProcessedAttachment(
            name=name,
            kind="unsupported",
            notes=(
                f"[Anhangstyp von '{name}' wird nicht unterstuetzt "
                f"({ctype or ext or 'unbekannt'}).]"
            ),
        )

    @staticmethod
    def _extract_hosted_content_ids(html_content: str) -> list[str]:
        if not html_content:
            return []
        ids = [unquote(match) for match in HOSTED_CONTENT_RE.findall(html_content)]
        # Reihenfolge beibehalten, Duplikate entfernen
        seen: set[str] = set()
        ordered: list[str] = []
        for item in ids:
            if item not in seen:
                seen.add(item)
                ordered.append(item)
        return ordered


def extract_document_text(name: str, data: bytes, content_type: str) -> str:
    """Extrahiert Klartext aus unterstützten Dokumentformaten."""
    lower = name.lower()
    ctype = content_type.lower()

    if lower.endswith(".pdf") or ctype == "application/pdf":
        return _extract_pdf(data)
    if lower.endswith(".docx") or "wordprocessingml" in ctype:
        return _extract_docx(data)

    text_exts = (".txt", ".md", ".csv", ".json", ".xml", ".log", ".html", ".htm")
    if (
        lower.endswith(text_exts)
        or ctype.startswith("text/")
        or ctype in {"application/json", "application/xml"}
    ):
        return _decode_text(data)

    return ""


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace").strip()


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        logger.warning("pypdf_not_installed")
        return ""

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text.strip():
            parts.append(page_text.strip())
    return "\n\n".join(parts).strip()


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        logger.warning("python_docx_not_installed")
        return ""

    document = Document(io.BytesIO(data))
    parts = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()
